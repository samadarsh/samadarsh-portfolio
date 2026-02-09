"""
Document loader for multiple file formats (PDF, TXT, DOCX).
Extracts text and metadata from uploaded documents.
"""

from pathlib import Path
from typing import Dict, List, Optional
from dataclasses import dataclass
import PyPDF2
import pdfplumber
from docx import Document
from src.utils.logger import get_logger

logger = get_logger()


@dataclass
class DocumentContent:
    """Container for document content and metadata."""
    text: str
    metadata: Dict[str, any]
    source: str
    page_count: Optional[int] = None


class DocumentLoader:
    """
    Multi-format document loader supporting PDF, TXT, and DOCX files.
    
    Features:
    - Automatic format detection based on file extension
    - Metadata extraction (filename, page numbers, etc.)
    - Error handling for corrupted or invalid files
    - Page-level text extraction for PDFs
    """
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.txt', '.docx'}
    
    def load_document(self, file_path: Path) -> DocumentContent:
        """
        Load document based on file extension.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            DocumentContent with extracted text and metadata
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
            Exception: For other loading errors
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {extension}")
        
        logger.info(f"Loading document: {file_path.name} (format: {extension})")
        
        try:
            if extension == '.pdf':
                return self.load_pdf(file_path)
            elif extension == '.txt':
                return self.load_txt(file_path)
            elif extension == '.docx':
                return self.load_docx(file_path)
        except Exception as e:
            logger.error(f"Error loading document {file_path.name}: {str(e)}")
            raise
    
    def load_pdf(self, file_path: Path) -> DocumentContent:
        """
        Load PDF document using pdfplumber for better text extraction.
        Falls back to PyPDF2 if pdfplumber fails.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            DocumentContent with extracted text and metadata
        """
        try:
            # Try pdfplumber first (better text extraction)
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                page_count = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text()
                    if page_text:
                        # Add page marker for better context tracking
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
                
                text = "\n\n".join(text_parts)
                
                metadata = {
                    "source": file_path.name,
                    "file_type": "pdf",
                    "page_count": page_count,
                    "loader": "pdfplumber"
                }
                
                logger.info(f"Successfully loaded PDF: {file_path.name} ({page_count} pages)")
                return DocumentContent(text=text, metadata=metadata, source=file_path.name, page_count=page_count)
                
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path.name}, trying PyPDF2: {str(e)}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_parts = []
                    page_count = len(pdf_reader.pages)
                    
                    for page_num, page in enumerate(pdf_reader.pages, start=1):
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"[Page {page_num}]\n{page_text}")
                    
                    text = "\n\n".join(text_parts)
                    
                    metadata = {
                        "source": file_path.name,
                        "file_type": "pdf",
                        "page_count": page_count,
                        "loader": "PyPDF2"
                    }
                    
                    logger.info(f"Successfully loaded PDF with PyPDF2: {file_path.name} ({page_count} pages)")
                    return DocumentContent(text=text, metadata=metadata, source=file_path.name, page_count=page_count)
                    
            except Exception as fallback_error:
                logger.error(f"Both PDF loaders failed for {file_path.name}: {str(fallback_error)}")
                raise Exception(f"Failed to load PDF: {str(fallback_error)}")
    
    def load_txt(self, file_path: Path) -> DocumentContent:
        """
        Load plain text file.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            DocumentContent with extracted text and metadata
        """
        try:
            # Try UTF-8 first, fall back to latin-1 if needed
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            except UnicodeDecodeError:
                logger.warning(f"UTF-8 decoding failed for {file_path.name}, trying latin-1")
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
            
            metadata = {
                "source": file_path.name,
                "file_type": "txt",
                "character_count": len(text)
            }
            
            logger.info(f"Successfully loaded TXT: {file_path.name} ({len(text)} characters)")
            return DocumentContent(text=text, metadata=metadata, source=file_path.name)
            
        except Exception as e:
            logger.error(f"Error loading TXT file {file_path.name}: {str(e)}")
            raise
    
    def load_docx(self, file_path: Path) -> DocumentContent:
        """
        Load DOCX (Microsoft Word) document.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            DocumentContent with extracted text and metadata
        """
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n\n".join(paragraphs)
            
            # Extract text from tables if any
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        table_texts.append(row_text)
            
            if table_texts:
                text += "\n\n[Tables]\n" + "\n".join(table_texts)
            
            metadata = {
                "source": file_path.name,
                "file_type": "docx",
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables)
            }
            
            logger.info(f"Successfully loaded DOCX: {file_path.name} ({len(paragraphs)} paragraphs)")
            return DocumentContent(text=text, metadata=metadata, source=file_path.name)
            
        except Exception as e:
            logger.error(f"Error loading DOCX file {file_path.name}: {str(e)}")
            raise
